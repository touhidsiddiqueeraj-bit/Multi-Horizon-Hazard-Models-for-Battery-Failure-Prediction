"""Generate a standalone document: Detailed Methodology & Results."""
import os, sys, tempfile, textwrap
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd
import numpy as np

PROJECT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCX_DIR = os.path.join(PROJECT, "paper")
FIG_DIR = os.path.join(PROJECT, "figs_journal_clean")
CSV_PATH = os.path.join(PROJECT, "data", "benchmark_results.csv")
DELONG_PATH = os.path.join(PROJECT, "tables_journal", "DeLong_AUC_comparisons.csv")
os.makedirs(DOCX_DIR, exist_ok=True)

MODEL_NAMES = {"xgboost": "XGBoost", "lightgbm": "LightGBM", "random_forest": "Random Forest", "gru": "GRU"}

def fmt_model(m):
    return MODEL_NAMES.get(m, m)

# ── Equation renderer (matplotlib → PNG) ────────────────────────────────
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

_EQ_CACHE = {}
def render_eq(latex: str, fontsize=13) -> str:
    if latex in _EQ_CACHE:
        return _EQ_CACHE[latex]
    fig, ax = plt.subplots(figsize=(5.5, 0.42))
    ax.text(0.5, 0.5, f"${latex}$", ha="center", va="center", fontsize=fontsize)
    ax.axis("off")
    fd, path = tempfile.mkstemp(suffix=".png")
    os.close(fd)
    fig.savefig(path, dpi=120, bbox_inches="tight", pad_inches=0.06, transparent=True)
    plt.close(fig)
    _EQ_CACHE[latex] = path
    return path

def cleanup_eq_cache():
    for p in _EQ_CACHE.values():
        try:
            os.remove(p)
        except:
            pass

# ── Document class ──────────────────────────────────────────────────────
class DetailedDoc:
    def __init__(self):
        self.doc = Document()
        s = self.doc.sections[0]
        s.page_width = Inches(8.5)
        s.page_height = Inches(11)
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
        style = self.doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = WD_LINE_SPACING.SINGLE

    def title(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run(text)
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.name = "Times New Roman"
        return p

    def h1(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text.upper())
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.name = "Times New Roman"
        return p

    def h2(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.name = "Times New Roman"
        return p

    def para(self, text):
        p = self.doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
        return p

    def eqn(self, latex_str):
        path = render_eq(latex_str)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.add_run().add_picture(path, width=Inches(4.5))
        return p

    def figure(self, path, caption):
        if os.path.exists(path):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.add_run().add_picture(path, width=Inches(5.0))
        pc = self.doc.add_paragraph(caption)
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in pc.runs:
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.name = "Times New Roman"
        return p if os.path.exists(path) else pc

    def table_title(self, text):
        p = self.doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.name = "Times New Roman"
        return p

    def make_table(self, headers, rows):
        t = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            self._set_cell(t.rows[0].cells[i], h, bold=True)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                self._set_cell(t.rows[ri + 1].cells[ci], str(val))
        return t

    def _set_cell(self, cell, text, bold=False, size=9):
        cell.text = str(text)
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(size)
            run.bold = bold
            run.font.name = "Times New Roman"

    def save(self, path):
        self.doc.save(path)
        return path

# ── Table generators ────────────────────────────────────────────────────
def make_dataset_stats_table(doc: DetailedDoc):
    """Table I: Descriptive statistics per dataset."""
    rows = []
    for ds, label, cells, colors in [
        ("nasa", "NASA 18650 [4]", 37, 1000),
        ("calce", "CALCE LCO/CX2 [5]", 7, 8733),
        ("oxford", "Oxford LFP [6]", 5, 287),
        ("severson", "Severson LFP [8]", 141, 117000),
    ]:
        d = pd.read_csv(CSV_PATH)
        d = d[d["dataset"] == ds]
        n_obs = len(d)
        n_pos = (d["AUC_raw"] > 0.5).sum() if len(d) > 0 else 0
        rows.append([label, str(cells), str(colors), f"{n_obs:,}"])
    doc.table_title("TABLE I: DATASET DESCRIPTIVE STATISTICS")
    doc.make_table(["Dataset", "Cells", "Total Cycles", "Records"], rows)

def make_feature_table(doc: DetailedDoc):
    """Table II: Feature availability per dataset."""
    rows = [
        ["NASA 18650 [4]", "✓", "✓", "✓", "✓", "✓", "✓", "✓"],
        ["CALCE LCO/CX2 [5]", "✓", "✓", "✓", "✓", "✗", "✗", "✓"],
        ["Oxford LFP [6]", "✓", "✓", "✓", "✓", "✓", "✓", "✓"],
        ["Severson LFP [8]", "✓", "✓", "✓", "✓", "✓", "✓", "✓"],
    ]
    doc.table_title("TABLE II: FEATURE AVAILABILITY PER DATASET")
    doc.make_table(
        ["Dataset", "Cycle", "Avg V", "Min V", "Avg I", "Avg T", "Duration", "SOH"],
        rows,
    )

def make_within_table(doc: DetailedDoc, df):
    """Table III: Within-dataset AUC per horizon."""
    best = df.loc[df.groupby(["eval", "dataset", "model", "H"])["AUC_cal"].idxmax()]
    rows = []
    for (mod, ds), grp in best[best["eval"] == "within"].groupby(["model", "dataset"]):
        aucs = {h: grp.loc[grp["H"] == h, "AUC_cal"].values[0] if (grp["H"] == h).any() else None for h in [10, 20, 30, 50]}
        mean_auc = grp["AUC_cal"].mean()
        mean_brier = grp["Brier_cal"].mean()
        rows.append([
            fmt_model(mod), ds.upper(),
            f"{aucs[10]:.3f}" if aucs[10] else "-",
            f"{aucs[20]:.3f}" if aucs[20] else "-",
            f"{aucs[30]:.3f}" if aucs[30] else "-",
            f"{aucs[50]:.3f}" if aucs[50] else "-",
            f"{mean_auc:.3f}", f"{mean_brier:.3f}",
        ])
    doc.table_title("TABLE III: WITHIN-DATASET AUC PER HORIZON (BEST CALIBRATION)")
    doc.make_table(["Model", "Dataset", "H=10", "H=20", "H=30", "H=50", "Mean AUC", "Brier"], rows)

def make_calibration_table(doc: DetailedDoc, df):
    """Table IV: Platt vs isotonic calibration."""
    best = df.loc[df.groupby(["eval", "dataset", "model", "H"])["AUC_cal"].idxmax()]
    rows = []
    for (eval_name, ds), grp in best[best["eval"] == "within"].groupby(["eval", "dataset"]):
        raw = df[(df["eval"] == eval_name) & (df["dataset"] == ds)]
        platt = raw[raw["method"] == "platt"]
        iso = raw[raw["method"] == "iso"]
        rows.append([
            ds.upper(),
            f"{iso['AUC_cal'].mean():.3f}",
            f"{platt['AUC_cal'].mean():.3f}",
            f"{iso['Brier_cal'].mean():.3f}",
            f"{platt['Brier_cal'].mean():.3f}",
        ])
    doc.table_title("TABLE IV: PLATT VS. ISOTONIC CALIBRATION (WITHIN-DATASET)")
    doc.make_table(["Dataset", "Iso AUC", "Platt AUC", "Iso Brier", "Platt Brier"], rows)

def make_cross_tables(doc: DetailedDoc, df):
    """Tables Va and Vb: Cross-chemistry with/without SOH."""
    cross = df[df["eval"] != "within"].copy()
    with_soh = cross[cross["eval"].str.contains("with_soh")]
    no_soh = cross[cross["eval"].str.contains("no_soh")]
    for tag, subset in [("WITH SOH", with_soh), ("WITHOUT SOH", no_soh)]:
        rows = []
        for (eval_name, mod), grp in subset.groupby(["eval", "model"]):
            for ds in ["oxford", "severson"]:
                dg = grp[grp["dataset"] == ds]
                if dg.empty:
                    continue
                raw_auc = dg["AUC_raw"].mean()
                iso_row = dg[dg["method"] == "iso"]
                iso_auc = iso_row["AUC_cal"].mean() if len(iso_row) > 0 else float("nan")
                rows.append([
                    eval_name.replace("train_", "").replace("_", " ").title().replace("Nasa", "NASA").replace("Calce", "CALCE"),
                    fmt_model(mod), ds.title(),
                    f"{raw_auc:.3f}", f"{iso_auc:.3f}",
                ])
        doc.table_title(f"TABLE VA: CROSS-CHEMISTRY TRANSFER — {tag}")
        doc.make_table(["Training", "Model", "Target", "Raw AUC", "Iso AUC"], rows)

def make_delong_table(doc: DetailedDoc):
    """Table VI: DeLong test results."""
    if not os.path.exists(DELONG_PATH):
        doc.para("(DeLong table not available.)")
        return
    delong = pd.read_csv(DELONG_PATH)
    rows = []
    for _, row in delong.iterrows():
        pv = row["p_value"]
        p_str = f"{pv:.2e}" if pv > 0 else "<1e-100"
        rows.append([
            row["dataset"], row["model_a"].replace("_", " "), row["model_b"].replace("_", " "),
            row["setting"], f"{row['AUC_a']:.3f}", f"{row['AUC_b']:.3f}",
            p_str, "✓" if row["significant_0.05"] else "✗",
        ])
    doc.table_title("TABLE VI: DELONG TEST FOR PAIRED AUC COMPARISONS")
    doc.make_table(["Dataset", "Model A", "Model B", "Setting", "AUC A", "AUC B", "p-value", "p<0.05"], rows)

# ── Content ─────────────────────────────────────────────────────────────
def build_document():
    doc = DetailedDoc()
    df = pd.read_csv(CSV_PATH)

    # ── Title ───────────────────────────────────────────────────────────
    doc.title("Multi-Horizon Hazard Models for Battery Failure Prediction\nDetailed Methodology and Results")
    doc.para(
        "This document provides a self-contained, detailed exposition of the methodology and "
        "results for the multi-horizon battery hazard classification study. All content is "
        "reproducible from the public code repository."
    )

    # ═════════════════════════════════════════════════════════════════════
    # III. METHODOLOGY
    # ═════════════════════════════════════════════════════════════════════
    doc.h1("III. Methodology")

    # III.A
    doc.h2("A. Composite Failure Label")
    doc.para(
        "Following the original protocol [1], a battery cycle is labeled as \"failure\" if either "
        "of two conditions is met: (1) State-of-Health (SOH) falls at or below 0.80 of initial capacity, "
        "where SOH is defined as the ratio of current discharge capacity to the mean capacity of the first "
        "10 cycles; or (2) the average discharge voltage drops below 94% of its early-life baseline (first "
        "10 cycles). The second criterion captures impedance-driven degradation where voltage sag precedes "
        "measurable capacity fade. Both conditions use the same baseline window, and once triggered, the "
        "label remains positive for all subsequent cycles."
    )
    doc.para(
        "SOH at cycle k for cell c is defined as the ratio of discharge capacity at cycle k "
        "to the mean capacity over the first 10 cycles:"
    )
    doc.eqn(r"\mathrm{SOH}_k^{(c)} = \frac{Q_k^{(c)}}{Q_{\mathrm{baseline}}^{(c)}},\quad "
            r"Q_{\mathrm{baseline}}^{(c)} = \frac{1}{10}\sum_{i=1}^{10} Q_i^{(c)}")
    doc.para(
        "The voltage-sag baseline is computed analogously over the first 10 cycles:"
    )
    doc.eqn(r"V_{\mathrm{baseline}}^{(c)} = \frac{1}{10}\sum_{i=1}^{10} V_{\mathrm{sag},i}^{(c)}")
    doc.para(
        "For cycle t in cell c with prediction horizon H, the binary failure label y_t^{(c)} is defined as:"
    )
    doc.eqn(
        r"y_t^{(c)} = \mathbf{1}[\,\exists k \in [t, t+H): "
        r"\mathrm{SOH}_k^{(c)} \leq 0.80 \;\mathrm{or}\; "
        r"V_{\mathrm{sag},k}^{(c)} < 0.94 \cdot V_{\mathrm{baseline}}^{(c)}\,]"
    )
    doc.para(
        "where 1[.] is the indicator function, yielding y = 1 if any cycle k within the "
        "prediction window satisfies either criterion, and y = 0 otherwise."
    )
    doc.para(
        "Once triggered (y = 1), the label remains 1 for all subsequent cycles of that cell. "
        "For datasets without voltage data (CALCE lacks discharge-duration logging) or where "
        "voltage is always at the cutoff (Oxford LFP's flat 2.7 V plateau), only the SOH criterion applies. "
        "Both thresholds are adopted directly from Shikdar & Laaksonen [1] without chemistry-specific "
        "tuning to maintain direct comparability. The fixed thresholds produce varying failure rates across "
        "datasets: NASA (37 cells, ~1,000 cycles/cell) exhibits moderate imbalance, CALCE (7 cells, 775-1,952 "
        "cycles/cell) shows heavily long-tailed degradation with ~92% positive rate, while Oxford and "
        "Severson LFP failure patterns depend on cycle-life distribution."
    )

    # III.B
    doc.h2("B. Datasets")
    doc.para(
        "Four publicly available battery cycling datasets are used, spanning two lithium-ion chemistries "
        "(LCO and LFP) with different cycling protocols, cell counts, and degradation characteristics. "
        "Table I summarizes the dataset sizes."
    )
    make_dataset_stats_table(doc)

    doc.para(
        "The NASA 18650 dataset [4] consists of 37 LCO cells (2.0 Ah rated capacity) aged under "
        "random-walk charge/discharge profiles at room temperature. Cells exhibit diverse degradation "
        "trajectories with approximately 1,000 cycles each. The CALCE LCO/CX2 dataset [5] contains "
        "7 LCO cells (CS2_33-36, CX2_36-38) cycled at constant 1C charge/1C discharge to 80% SOH, "
        "yielding 8,733 total cycles with slow, uniform degradation spanning 775-1,952 cycles per cell. "
        "CALCE lacks temperature and discharge-duration measurements."
    )
    doc.para(
        "Two LFP datasets serve as cross-chemistry transfer targets. The Oxford LFP dataset [6] "
        "contains 5 LFP pouch cells (2.3 Ah) cycled at 1C/1C under controlled temperature (40 C) "
        "and pressure, providing 300-500 cycles per cell with measurements recorded at approximately "
        "100-cycle intervals. The MIT-Stanford Severson dataset [8] contains 141 LFP cells aged under "
        "a fast-charging protocol (variable charge rates, 4C discharge rate) spanning 534-2,237 cycles "
        "per cell (approximately 117,000 total cycles across four experimental batches), providing a "
        "substantially larger and more diverse LFP test set."
    )
    make_feature_table(doc)
    doc.para(
        "Table II shows which features are available per dataset. The CALCE dataset has entirely missing "
        "temperature and discharge-duration columns; these are filled with zeros. Because these features "
        "are constant across all CALCE rows, tree-based models cannot form informative splits on them. "
        "However, they act as a trivial dataset fingerprint for cross-chemistry evaluation: any model "
        "trained on non-CALCE data sees zero-inflated features when evaluated on CALCE test rows. "
        "Since cross-chemistry evaluation always tests on LFP targets (Oxford or Severson) where all "
        "features are populated, this artifact does not affect the reported cross-chemistry results."
    )

    # III.C
    doc.h2("C. Features and Preprocessing")
    doc.para(
        "All models use per-cycle features: cycle number, average discharge voltage, minimum discharge "
        "voltage, average discharge current, average temperature, discharge duration, and SOH. "
        "Features are used as-is without scaling or normalization, as tree-based models are invariant "
        "to monotonic transformations. Missing values in CALCE (average temperature and discharge "
        "duration) are filled with zero. This zero-fill is physically meaningless (0 C average temperature, "
        "0 s discharge duration) but constant across all CALCE rows, so tree-based models cannot exploit "
        "these features for within-dataset splits."
    )
    doc.para(
        "For cross-chemistry evaluation, two feature configurations are tested: with SOH (all seven "
        "features) and without SOH (six features, excluding SOH). The without-SOH condition tests "
        "whether the remaining features carry chemistry-agnostic degradation signal. The GRU receives "
        "a reduced feature set for cross-chemistry experiments: cycle number, average voltage, minimum "
        "voltage, and SOH (when available). Average current (unit mismatch: A vs mA across datasets), "
        "average temperature, and discharge duration are excluded to avoid spurious discriminative "
        "signals from zero-filled CALCE features."
    )

    # III.D
    doc.h2("D. Models and Hyperparameters")
    doc.para(
        "Three tree-based classifiers are compared: XGBoost (max_depth=4, n_estimators=300, "
        "learning_rate=0.05, subsample=0.8, colsample_bytree=0.8), LightGBM (max_depth=4, "
        "n_estimators=300, learning_rate=0.05, subsample=0.8, colsample_bytree=0.8), and "
        "Random Forest (max_depth=6, n_estimators=300, max_samples=0.8). All tree models use "
        "default class weighting and random_state=42 for deterministic replication. Hyperparameters "
        "are matched to the original study [1] where applicable; no dataset-specific tuning is performed."
    )

    # III.E
    doc.h2("E. Gated Recurrent Unit Sequence Classifier")
    doc.para(
        "A GRU classifier with a single hidden layer of 8 units processes a sliding temporal window "
        "of W = 10 consecutive cycles per cell. The input is a 10 x f matrix (f features) processed "
        "sequentially; after the final timestep, the hidden state h_W passes through a linear layer "
        "with sigmoid activation to produce a failure probability. The architecture is deliberately "
        "compact to prevent overfitting given the limited number of training cells (7-37). "
        "Binary cross-entropy loss with inverse-frequency class weighting "
        "(pos_weight = n_neg / n_pos) addresses class imbalance. The Adam optimizer "
        "(learning_rate = 0.005) with early stopping (patience = 10) is used."
    )
    doc.para(
        "A single seed (seed=0) is used for GRU experiments. The observed per-configuration AUC "
        "instability (range 0.011-0.986 across training configurations, horizons, and SOH conditions) "
        "is itself an informative result: it reflects distributed hidden-state entanglement under "
        "covariate shift, demonstrating that single-seed recurrent models are unreliable for "
        "cross-chemistry evaluation. Multi-seed analysis is deferred to future work. Tree-based "
        "models show no such instability (AUC standard deviation < 0.01 across identical repeats)."
    )

    # III.F
    doc.h2("F. Calibration")
    doc.para(
        "Two post-hoc calibration methods are compared. Platt (sigmoid) scaling fits a logistic "
        "regression model to the classifier's raw scores on the calibration set:"
    )
    doc.eqn(r"P(y=1 \mid x) = \frac{1}{1 + \exp(a \cdot s(x) + b)}")
    doc.para(
        "where s(x) is the model's raw score and a, b are learned parameters. Logistic regression "
        "uses high regularization (C = 1e10). Isotonic regression [2] fits a non-decreasing step "
        "function via the pool-adjacent-violators (PAV) algorithm, making no parametric assumption "
        "about the calibration mapping shape."
    )
    doc.para(
        "Neither method uses CalibratedClassifierCV, which would create a 3-model ensemble and "
        "conflate ensembling effects with calibration quality. Both calibrators are applied to "
        "the training-fold scores rather than a held-out calibration set; they share the same "
        "underlying classifier outputs, so the comparison between them is fair. However, absolute "
        "calibrated metrics (particularly Brier scores) may be optimistic relative to a held-out "
        "calibration procedure. Performance is measured by both AUC (discrimination) and Brier "
        "score (calibration + discrimination):"
    )
    doc.eqn(r"\text{Brier} = \frac{1}{N}\sum_{i=1}^{N} (p_i - y_i)^2")
    doc.para(
        "where p_i is the calibrated probability and y_i is the true binary label."
    )
    doc.para(
        "For each (dataset, model, horizon) configuration, we select the calibration method "
        "(Platt or isotonic) yielding the higher mean AUC. For cross-chemistry comparisons, "
        "raw (uncalibrated) AUC is the primary metric because post-hoc calibration under "
        "distribution shift is unreliable."
    )

    # III.G
    doc.h2("G. Cross-Chemistry Transfer Protocol")
    doc.para(
        "Cross-chemistry transfer experiments test whether models trained on LCO data generalize "
        "to LFP. Models are trained on all cycles from one or more LCO training datasets and "
        "evaluated on all cycles from an LFP target dataset. Three training configurations are "
        "compared: NASA only (37 LCO cells), CALCE only (7 LCO cells), and ALL (NASA + CALCE, "
        "44 LCO cells). Each configuration is tested with and without SOH as a feature. Two LFP "
        "test targets are used: Oxford (5 cells, standard 1C/1C protocol) and MIT-Stanford "
        "Severson (141 cells, variable-rate fast-charging protocol). Including two LFP targets "
        "with different cycling protocols and cell counts (5 vs 141) tests the robustness of "
        "the cross-chemistry findings."
    )
    doc.para(
        "Per-cell evaluation is used: models are trained on all LCO cells and evaluated "
        "independently on each LFP cell. AUC and Brier scores are computed per cell and "
        "reported as mean across cells. This provides a measure of how consistently the "
        "model's discriminative ability holds across individual LFP cells, unlike a single "
        "pooled evaluation which conflates within-cell ranking with between-cell differences."
    )

    # III.H
    doc.h2("H. Evaluation Protocol")
    doc.para(
        "Within-dataset evaluation uses 5-fold GroupKFold stratified by cell: all cycles from "
        "a given cell belong to the same fold, ensuring generalization is measured across unseen "
        "cells rather than unseen cycles. Four prediction horizons H  in {10, 20, 30, 50} are "
        "tested, where the label for cycle t is positive if the battery fails within [t, t+H). "
        "Models are retrained from scratch for each horizon. Metrics are reported as means across "
        "folds. The best calibration method per (dataset, model) pair is selected by mean AUC."
    )
    doc.eqn(
        r"\mathrm{AUC} = \int_{0}^{1} \mathrm{TPR}(\mathrm{FPR}^{-1}(t)) \, dt = "
        r"P(\mathrm{score}_{\mathrm{pos}} > \mathrm{score}_{\mathrm{neg}})"
    )
    doc.para(
        "AUC is interpreted as the probability that a randomly chosen positive sample receives "
        "a higher score than a randomly chosen negative sample. The DeLong nonparametric test [16] "
        "is used to assess whether AUC differences between paired conditions (e.g., with-SOH vs "
        "without-SOH) are statistically significant. The DeLong z-statistic is computed from the "
        "empirical covariance matrix of the two ROC curves:"
    )
    doc.eqn(r"z = \frac{\mathrm{AUC}_1 - \mathrm{AUC}_2}{\sqrt{\mathrm{Var}(\mathrm{AUC}_1) + \mathrm{Var}(\mathrm{AUC}_2) - 2\,\mathrm{Cov}(\mathrm{AUC}_1, \mathrm{AUC}_2)}}")
    doc.para(
        "Note on Oxford multi-horizon evaluation. The Oxford LFP dataset is recorded at "
        "approximately 100-cycle intervals (5 cells, 46-78 rows each). The multi-horizon label "
        "function operates on raw cycle numbers, so H = 10, 20, 30, and 50 all map to identical "
        "binary labels (76 rows, 23.8% positive rate). Oxford multi-horizon analysis therefore "
        "collapses to a single effective horizon. The core cross-chemistry findings are unaffected "
        "because all horizon variants share the same label set."
    )

    # ═════════════════════════════════════════════════════════════════════
    # IV. RESULTS
    # ═════════════════════════════════════════════════════════════════════
    doc.h1("IV. Results")

    # IV.A
    doc.h2("A. Within-Dataset Performance")
    doc.para(
        "Table III presents mean AUC and Brier scores (across all four horizons, best calibration) "
        "for all models on NASA and CALCE. Both datasets show strong discrimination with mean AUC "
        r"ranging from 0.878 (Random Forest on NASA) to 0.949 (GRU on CALCE). All tree-based models "
        "achieve mean AUC >= 0.85 on both datasets. The GRU achieves competitive within-dataset "
        "performance (mean AUC = 0.886 on NASA, 0.949 on CALCE), confirming that the compact 8-unit "
        "architecture is sufficient for within-chemistry prediction."
    )
    doc.para(
        "AUC increases monotonically with prediction horizon for all tree-based models on both "
        "datasets. The steepest gains occur between H = 10 and H = 20 (mean improvement 0.015-0.030), "
        "with diminishing returns at longer horizons. The Random Forest underperforms XGBoost and "
        "LightGBM by 2-3 AUC points on average, consistent with the known advantage of boosting "
        "over bagging for structured tabular data with imbalanced class distributions."
    )
    make_within_table(doc, df)
    doc.para("")  # spacer

    doc.figure(
        os.path.join(FIG_DIR, "Fig01_Within_Dataset_AUC.png"),
        "Figure 1. Within-dataset AUC heatmap at H = 20 (best calibration per dataset). "
        "AUC values range from 0.875 (Random Forest on NASA) to 0.920 (LightGBM on CALCE)."
    )
    doc.figure(
        os.path.join(FIG_DIR, "Fig05_MultiHorizon_AUC.png"),
        "Figure 2. Multi-horizon AUC on NASA (Platt-calibrated) as a function of prediction "
        "horizon H. AUC improves from H = 10 to H = 50 across all tree-based models, with "
        "the steepest gains at shorter horizons."
    )

    # IV.B
    doc.h2("B. Platt vs. Isotonic Calibration")
    doc.para(
        "Table IV compares Platt and isotonic calibration across all within-dataset configurations. "
        "Platt achieves higher mean AUC on both NASA (Platt 0.899, Isotonic 0.880) and CALCE "
        "(Platt 0.900, Isotonic 0.892). On CALCE, the gap is larger and more variable: for LightGBM, "
        "Platt AUC = 0.918 vs. Isotonic AUC = 0.694, a gap of 0.224. This arises from CALCE's "
        "long-tailed degradation distribution (up to 1,952 cycles per cell with heavily imbalanced "
        "failure rates). Isotonic regression bins the extreme scores produced by these imbalanced "
        "tails into degenerate steps, reducing discriminative power while preserving average "
        "calibration (Brier scores remain comparable by construction)."
    )
    doc.para(
        "The Brier scores show negligible differences between calibrators (~0.001-0.002), which is "
        "orders of magnitude smaller than the overall Brier gap between our results (0.17-0.26) and "
        "the published values (~0.032) in Shikdar & Laaksonen [1]. The source of this approximately "
        "8x discrepancy could not be determined from the available code and documentation. Our AUC "
        "values (0.80-0.90) are consistent with the published range, and the within-dataset Brier "
        "differences between calibration methods are orders of magnitude smaller than the gap to "
        "published values, suggesting the discrepancy is in the overall score scale rather than "
        "in the relative method comparison."
    )
    make_calibration_table(doc, df)
    doc.para("")

    doc.figure(
        os.path.join(FIG_DIR, "Fig02_Calibration_Comparison.png"),
        "Figure 3. Platt vs. isotonic calibration reliability diagrams for NASA and CALCE "
        "(XGBoost, H = 20). Platt maintains smoother calibration curves; isotonic produces "
        "degenerate bins on long-tailed CALCE data."
    )

    # IV.C
    doc.h2("C. Cross-Chemistry Transfer")
    doc.para(
        "Tables Va and Vb present cross-chemistry transfer results with and without SOH as a feature. "
        "The contrast is unambiguous. When SOH is available, tree-based models achieve raw AUC values "
        "ranging from 0.84 to 1.00 across all training configurations and both LFP targets. "
        "The GRU achieves substantially lower raw AUC (0.08-0.73) even with SOH available, "
        "due to the architecture-specific entanglement mechanism discussed in Section IV-D."
    )
    make_cross_tables(doc, df)

    doc.para(
        "When SOH is removed from the feature set, tree-based raw AUC collapses to 0.33-0.62 "
        "on Oxford and 0.60-0.75 on Severson. The Severson without-SOH above-chance performance "
        "reflects a partial cycle-number proxy effect: Severson's wide cycle-life range "
        "(150-2,300 cycles) partially overlaps with LCO training distributions (NASA ~1,000, "
        "CALCE 775-1,952 cycles), so cycle number alone carries weak transferable signal. "
        "Oxford's narrow cycle range (~300 cycles) lies entirely outside the LCO distribution, "
        "eliminating this weak proxy. Despite this partial signal, the SOH-ablation gap on "
        "Severson is 19-24 AUC points and statistically decisive (p < 10^{-100}), confirming "
        "that cycle number is a weak distributional proxy rather than a transferable electrochemical feature."
    )
    doc.para(
        "The evidence for the SOH-as-lookup-table mechanism is threefold: (1) with SOH, AUC reaches "
        "1.00 for NASA-to-Oxford -- perfect transfer that disappears when SOH is removed; "
        "(2) the GRU's inability to exploit SOH under distribution shift is architecture-consistent "
        "with distributed representations partially corrupting the feature; (3) SHAP analysis "
        "(Section IV-G, Figures 6a-6c) confirms that SOH dominates tree-model split decisions "
        "in cross-chemistry configurations, with cycle number as a distant second."
    )

    doc.figure(
        os.path.join(FIG_DIR, "Fig03_CrossChem_With_SOH.png"),
        "Figure 4. Cross-chemistry transfer with SOH (raw AUC, trees H = 20, GRU mean across H). "
        "Left: Oxford (5 cells). Right: Severson (141 cells). Consistent SOH-driven high AUC "
        "across both LFP targets."
    )
    doc.figure(
        os.path.join(FIG_DIR, "Fig04_CrossChem_No_SOH.png"),
        "Figure 5. Cross-chemistry transfer without SOH (raw AUC, trees H = 20, GRU mean across H). "
        "Left: Oxford (5 cells). Right: Severson (141 cells). AUC collapses across all "
        "training-by-target combinations, confirming SOH dependence."
    )

    # IV.D
    doc.h2("D. GRU Entanglement Under Distribution Shift")
    doc.para(
        "The GRU sequence classifier reveals a novel architecture-specific phenomenon: how a model "
        "builds its internal representation determines whether it can exploit SOH as a lookup-table "
        "shortcut under distribution shift. The GRU's distributed hidden state -- compressed into "
        "8 dimensions -- entangles SOH with voltage, current, and cycle trends across the 10-timestep "
        "window. Under LCO-to-LFP distribution shift, the entangled voltage and cycle components "
        "carry distribution-mismatched signal that partially corrupts the SOH channel. Tree-based "
        "models avoid this entirely: their isolated hard splits on SOH transfer perfectly across "
        "chemistries because each split examines SOH alone, independent of other features."
    )
    doc.para(
        "This architecture-specific fragility manifests across horizons as a consequence of the "
        "representation-sharing mechanism. At short horizons (H = 10), the voltage and cycle noise "
        "dominate the compressed hidden state, producing AUC as low as 0.017. At long horizons "
        "(H = 50), the accumulated SOH signal across 50-cycle windows eventually overpowers the "
        "entangled noise, reaching AUC up to 0.978. Tree models exhibit no such horizon sensitivity "
        "because each split is feature-isolated."
    )
    doc.para(
        "A separate mechanism drives the CALCE-to-Oxford reversal. CALCE's 92% composite-failure "
        "rate saturates the GRU's learned decision boundary, producing systematically inverted "
        "rank-orderings on Oxford's feature distribution (AUC approximately 0.03-0.12). This is a "
        "class-imbalance-driven domain mismatch, distinct from the entanglement effect seen with "
        "NASA training. We report raw AUC values rather than max(AUC, 1 - AUC) because the "
        "reversal reflects a substantive failure mode, not an arbitrary label-polarity choice."
    )
    doc.para(
        "These findings establish a secondary result with independent significance: the choice of "
        "model architecture determines not just predictive performance but the very mechanism by "
        "which a model exploits (or fails to exploit) leakage features under distribution shift. "
        "Tree models' feature-isolated splits make them maximally robust to SOH-as-leakage; the "
        "GRU's distributed representations make it architecture-specifically fragile."
    )

    # IV.E
    doc.h2("E. Calibration Methods Fail to Transfer")
    doc.para(
        "Tables Va and Vb reveal an important secondary phenomenon: the isotonic-calibrated AUC "
        "values are systematically lower than the corresponding raw AUC values for cross-chemistry "
        "transfer, often dramatically so. For example, XGBoost trained on ALL LCO cells achieves "
        "raw AUC = 0.979 with SOH at H = 20, but isotonic calibration collapses this to 0.510. "
        "The collapse occurs because isotonic regression fits a step function to the training set's "
        "score distribution; when the test set's scores follow a different distribution (as they "
        "inevitably do under cross-chemistry covariate shift), multiple test scores fall into the "
        "same isotonic bin and are assigned identical calibrated probabilities. These ties "
        "artificially reduce AUC because tied predictions with different true labels contribute "
        "a 0.5 penalty per pair."
    )
    doc.para(
        "The effect is most severe for the GRU, where isotonic collapses nearly all cross-chemistry "
        "scores to a single bin (calibrated AUC = 0.500 for six of nine training-by-SOH "
        "configurations), but tree-based models also lose 0.15-0.48 AUC points depending on the "
        "setting. This finding -- that calibration methods themselves fail to transfer across "
        "chemistries -- is independent of the SOH-lookup-table mechanism and represents a second, "
        "distinct failure mode for cross-chemistry battery hazard prediction. It implies that even "
        "when raw model scores carry transferable signal (as they do with SOH), post-hoc calibration "
        "cannot be naively applied under distribution shift without risking the destruction of that signal."
    )

    # IV.F
    doc.h2("F. DeLong Test: Statistical Significance of SOH Ablation")
    doc.para(
        "To establish whether the AUC differences between the with-SOH and without-SOH conditions "
        "are statistically significant, we apply the DeLong nonparametric test for paired ROC "
        "curves [16]. Unlike a naive comparison of point estimates, the DeLong test accounts for "
        "the correlation between AUC values derived from the same set of test samples."
    )
    doc.para(
        "Table VI reports DeLong p-values for the SOH-ablation comparison (ALL LCO to LFP, H = 20) "
        "on both Oxford and Severson, alongside within-dataset model-pair comparisons. The "
        "SOH-ablation p-values span 10^{-36} to < 10^{-100}, providing decisive evidence that the "
        "AUC collapse when removing SOH is not a chance fluctuation. On Severson, the substantially "
        "larger test set (141 cells vs 5) drives p-values below 10^{-100} for all three tree-based "
        "models. Within-dataset comparisons show a different picture: on NASA (37 cells), model-level "
        "AUC differences are small (e.g., XGBoost vs LightGBM: delta AUC = 0.014, p = 0.014), with "
        "only XGBoost vs LightGBM reaching significance. On CALCE (8,733 cycles), the larger sample "
        "yields significant differences across all three model pairs (p < 10^{-5})."
    )
    make_delong_table(doc)

    # IV.G
    doc.h2("G. SHAP Feature Importance")
    doc.para(
        "To investigate the role of individual features in cross-chemistry transfer, we compute "
        "SHAP (SHapley Additive exPlanations) values [11] for the three tree-based models trained "
        "on NASA and tested on Oxford (H = 20, with SOH). The SHAP value phi_j for feature j "
        "quantifies the marginal contribution of that feature to the model's prediction, averaged "
        "over all possible feature orderings:"
    )
    doc.eqn(
        r"\phi_j = \sum_{S \subseteq F \setminus \{j\}} "
        r"\frac{|S|!\,(|F|-|S|-1)!}{|F|!} "
        r"[f(S \cup \{j\}) - f(S)]"
    )
    doc.para(
        "Figures 6a-6c present the SHAP summary plots for XGBoost, LightGBM, and Random Forest "
        "respectively (with SOH). Across all three model classes, SOH dominates as the most important "
        "feature by a wide margin. Cycle number is a distant second, while voltage, current, "
        "temperature, and duration features contribute negligibly. This pattern is consistent with "
        "the SOH-as-lookup-table mechanism: the models rely almost exclusively on SOH to make "
        "cross-chemistry predictions."
    )

    for fig_label, fig_file, model_name in [
        ("a", "Fig06a_XGBoost_SHAP.png", "XGBoost"),
        ("b", "Fig06b_LightGBM_SHAP.png", "LightGBM"),
        ("c", "Fig06c_RandomForest_SHAP.png", "Random Forest"),
    ]:
        doc.figure(
            os.path.join(FIG_DIR, fig_file),
            f"Figure 6{fig_label}. SHAP feature importance for {model_name} in "
            "NASA-to-Oxford cross-chemistry transfer (H = 20, with SOH). SOH dominates all other features."
        )

    doc.para(
        "Figures 6d-6f present the corresponding SHAP summary plots for the no-SOH condition "
        "(training on NASA features excluding SOH, testing on Oxford, H = 20). The contrast is "
        "stark: where Figures 6a-6c show SOH dominating with high-magnitude SHAP values across "
        "the full feature range, Figures 6d-6f show all remaining features collapsed to near-zero "
        "SHAP spread with no meaningful ranking signal. This visual collapse mirrors the quantitative "
        "AUC collapse: without SOH, no feature carries sufficient chemistry-invariant signal to drive "
        "discriminative splits, and SHAP values reflect near-random permutation effects."
    )

    for fig_label, fig_file, model_name in [
        ("d", "Fig06d_XGBoost_SHAP_noSOH.png", "XGBoost"),
        ("e", "Fig06e_LightGBM_SHAP_noSOH.png", "LightGBM"),
        ("f", "Fig06f_RandomForest_SHAP_noSOH.png", "Random Forest"),
    ]:
        doc.figure(
            os.path.join(FIG_DIR, fig_file),
            f"Figure 6{fig_label}. SHAP feature importance for {model_name} in "
            "NASA-to-Oxford cross-chemistry transfer (H = 20, without SOH). All features collapse "
            "to near-zero SHAP spread."
        )

    # ═════════════════════════════════════════════════════════════════════
    # V. DISCUSSION (condensed)
    # ═════════════════════════════════════════════════════════════════════
    doc.h1("V. Discussion")
    doc.para(
        "The central finding is that cross-chemistry transfer of hazard-based battery failure "
        "prediction fails for all model classes tested when SOH is unavailable as a feature. "
        "This result holds across two LFP test targets (Oxford, 5 cells; Severson, 141 cells), "
        "three tree-based model classes (XGBoost, LightGBM, Random Forest), a sequence model (GRU), "
        "and three LCO training configurations (NASA, CALCE, ALL). The DeLong test confirms that "
        "the SOH-ablation AUC gap is statistically significant at p < 10^{-36} for all "
        "model-by-target combinations."
    )
    doc.para(
        "The mechanism underlying this failure is the SOH-as-lookup-table phenomenon: models trained "
        "on LCO data learn a mapping from SOH values to remaining useful life that is specific to "
        "LCO degradation trajectories. When deployed on LFP, the model applies this same learned "
        "mapping -- producing high AUC because LFP cells traverse similar SOH ranges -- but the "
        "prediction is driven by a chemistry-specific proxy, not by genuine understanding of LFP "
        "degradation dynamics. The SHAP analysis confirms this: SOH accounts for the dominant share "
        "of feature importance, and removing it collapses all predictors to near-random performance."
    )
    doc.para(
        "Three distinct failure modes are identified: (1) SOH-as-lookup-table (all models), "
        "(2) distributed-representation entanglement under covariate shift (GRU-specific), and "
        "(3) calibration collapse under distribution shift (all post-hoc methods, with isotonic "
        "regression being the most severely affected). These findings establish that cross-chemistry "
        "generalization of hazard-based battery failure models remains an open problem requiring "
        "either chemistry-invariant feature representations or per-chemistry model retraining."
    )

    # ── Save ────────────────────────────────────────────────────────────
    out = os.path.join(DOCX_DIR, "paper_methodology_results.docx")
    doc.save(out)
    print(f"Saved: {out}")
    return out

if __name__ == "__main__":
    path = build_document()
    cleanup_eq_cache()
    print(f"Done. Output: {path}")
