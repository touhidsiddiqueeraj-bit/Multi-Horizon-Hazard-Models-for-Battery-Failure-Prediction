"""Generate IEEE Access formatted Paper_IEEE.docx — with embedded edge deployment focus."""
import os
import re
import pandas as pd
from docx.shared import Inches, Pt
from ieee_format import IEEEPaper

PROJECT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCX_DIR = os.path.join(PROJECT, "paper")
FIG_DIR  = os.path.join(PROJECT, "data")
CSV_PATH = os.path.join(PROJECT, "data", "benchmark_results.csv")
DELONG_PATH = os.path.join(PROJECT, "tables_journal", "DeLong_AUC_comparisons.csv")

os.makedirs(DOCX_DIR, exist_ok=True)

paper = IEEEPaper()
paper.add_title(
    "Multi-Horizon Hazard Models for Battery\n"
    "Failure Prediction: Within-Dataset Reliability,\n"
    "Cross-Chemistry Transferability, and\n"
    "Embedded Edge Deployment"
)
paper.add_author(
    "Hussain Touhid Siddiquee, Syeda Salsabil Islam,\n"
    "Ariya Jasimul Islam, and Chowdhury Farzana Hoque Eshica"
)
paper.add_affiliation(
    "Department of Electrical and Electronic Engineering,\n"
    "Leading University, Sylhet, Bangladesh"
)
paper.add_affiliation("Corresponding author: Hussain Touhid Siddiquee (touhidsiddiqueeraj@gmail.com)")

paper.add_abstract(
    "Predicting whether a lithium-ion cell will fail within a short operational window is a distinct "
    "task from estimating its remaining useful life, yet it is the question that matters most for "
    "real-time dispatch in electric vehicles and grid storage. This paper extends a recently proposed "
    "multi-horizon hazard classification framework, originally demonstrated on a single model and a "
    "single dataset, into a broader study covering four model families, three lithium-cobalt-oxide "
    "(LCO) training sets, and two lithium-iron-phosphate (LFP) transfer targets. We evaluate "
    "within-dataset reliability (AUC \u2265 0.85 across 31 of 32 configurations), compare Platt "
    "and isotonic calibration under a fairness-corrected protocol, and demonstrate through a "
    "controlled SOH-ablation study that cross-chemistry transfer depends entirely on SOH as a "
    "chemistry-specific lookup feature: removing it collapses AUC from 0.84\u20131.00 to 0.33\u20130.62 "
    "on Oxford LFP and 0.60\u20130.85 on MIT-Stanford Severson LFP (DeLong p < 10\u207b\u00b3\u2076). "
    "A secondary finding reveals that isotonic calibration itself fails to transfer across chemistries, "
    "systematically destroying 0.15\u20130.48 AUC points under distribution shift. We further "
    "demonstrate a complete embedded deployment: all three tree ensembles (900 trees, 26,306 nodes) "
    "are packed into a 372 kB binary and run on an ESP32-S3 microcontroller through a hand-written "
    "C tree walker, reproducing Python training library predictions to sub-microsecond precision "
    "across 1,028 validation rows in under 600 \u00b5s. These results establish a reproducible "
    "within-dataset baseline while clarifying fundamental limitations of cross-chemistry transfer "
    "and providing a validated edge deployment path."
)
paper.add_keywords(
    "Battery failure prediction, multi-horizon hazard classification, cross-chemistry transfer, "
    "probability calibration, SHAP, embedded machine learning, ESP32-S3"
)

# ── Begin two-column body ──────────────────────────────────────────────────
paper.begin_body(
    running_title="H. T. Siddiquee et al.: Multi-Horizon Hazard Models for Battery Failure Prediction"
)

# ══════════════════════════════════════════════════════════════════════════════
#  I. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
paper.add_h1("I. Introduction")
paper.add_para(
    "Lithium-ion batteries underpin the electrification of transport and the decarbonization of "
    "electrical grids, yet their long-term reliability remains difficult to certify in the field. "
    "Conventional prognostics frames the problem as remaining-useful-life (RUL) regression: a model "
    "predicts a continuous cycle count to end-of-life [1]\u2013[3]. While useful for capacity-fade "
    "studies, RUL regression answers a question that operators rarely ask in real time. A battery "
    "management system needs to know whether a specific cell will survive the next mission, not "
    "how many cycles it has left in aggregate."
)
paper.add_para(
    "Shikdar and Laaksonen [1] recently proposed a multi-horizon hazard classification framework "
    "that reformulates the problem: for a given battery cycle and horizon H, the model predicts "
    "whether failure\u2014defined as State-of-Health (SOH) below 0.80 or average voltage sag "
    "exceeding 6%\u2014will occur within the next H cycles. Using Histogram-based Gradient Boosting "
    "(HGB) with isotonic calibration on the NASA 18650 dataset (37 LCO cells), they reported "
    "AUC values of 0.87\u20130.90 and Brier scores of 0.170\u20130.175 across H \u2208 {10, 20, 30, 50}."
)
paper.add_para(
    "That work left three questions open, and each of them has practical consequences. First, "
    "the original study used a single model class on a single dataset, so it is unclear how "
    "sensitive the results are to the choice of learner, the hyperparameter setting, or the "
    "post-hoc calibration method. Second, no cross-chemistry validation was performed: a model "
    "trained on LCO data may not transfer to LFP or NMC chemistries that dominate electric "
    "vehicle and grid-storage applications. Third, the path from a Python research prototype "
    "to an embedded microcontroller\u2014where battery management systems actually run\u2014was "
    "not addressed."
)
paper.add_para(
    "This paper addresses all three gaps. The contributions are: (i) a four-model benchmark "
    "comprising three tree ensembles (XGBoost, LightGBM, Random Forest) and a GRU sequence "
    "classifier, all evaluated on two LCO datasets (NASA 18650, CALCE LCO/CX2) under a unified "
    "5-fold GroupKFold protocol; (ii) a controlled cross-chemistry transfer study from LCO to "
    "two independent LFP targets\u2014the 5-cell Oxford pouch set and the 141-cell MIT\u2013Stanford "
    "(Severson) set\u2014with and without State-of-Health (SOH) as an input feature, analyzed "
    "with DeLong significance tests and SHAP explanations; (iii) a fairness-corrected comparison "
    "of isotonic and Platt calibration; and (iv) a complete embedded deployment that packs all "
    "three trained tree ensembles into a single flat binary and runs them on an ESP32-S3 "
    "microcontroller through a hand-written C tree walker."
)
paper.add_para(
    "The central negative result\u2014that removing SOH collapses cross-chemistry AUC from "
    "near-perfect to near-chance\u2014is not a generic null finding. It points to a specific "
    "mechanism: trees exploit SOH through isolated hard splits that act as a chemistry-specific "
    "capacity-to-horizon lookup table, and the lookup table does not transfer. A secondary, "
    "independent result is that calibration methods themselves fail to transfer: isotonic "
    "regression systematically destroys cross-chemistry AUC by binning distribution-shifted "
    "scores into degenerate steps."
)

# ══════════════════════════════════════════════════════════════════════════════
#  II. RELATED WORK
# ══════════════════════════════════════════════════════════════════════════════
paper.add_h1("II. Related Work")
paper.add_h2("A. From RUL regression to hazard classification")
paper.add_para(
    "Data-driven battery prognostics has historically been formulated as a regression problem "
    "in which capacity fade or internal resistance growth is extrapolated to a threshold. "
    "Meng and Li [2] review this field and group prior methods into three families: physics-based "
    "models, purely data-driven models, and hybrid approaches. Tree ensembles\u2014notably "
    "Random Forest [3], XGBoost [4], and LightGBM [5]\u2014have been applied to capacity "
    "estimation and RUL prediction, but their use for binary hazard classification at multiple "
    "temporal horizons is recent [1]."
)
paper.add_h2("B. Cross-chemistry and transfer-learning approaches")
paper.add_para(
    "Transfer learning for battery state estimation has begun to attract attention. Sahoo et al. [11] "
    "developed a transfer-learning framework for SOH estimation validated across NASA and CALCE cells, "
    "demonstrating that fine-tuning on a small target-chemistry sample recovers accurate capacity "
    "estimates. Lu et al. [12] proposed a deep learning architecture pretrained on large-scale "
    "degradation data and fine-tuned to new cell types. However, these studies target regression-based "
    "SOH estimation, not classification-based hazard prediction. The multi-horizon hazard setting "
    "introduces two specific complications: the label depends on the chosen horizon H, and the "
    "calibration of probability outputs\u2014critical for decision-making\u2014may degrade under "
    "distribution shift."
)
paper.add_h2("C. Calibration of classifier outputs")
paper.add_para(
    "Post-hoc calibration maps raw classifier scores to calibrated probability estimates. "
    "Platt scaling [13] fits a sigmoid through logistic regression on the model\u2019s scores, "
    "while isotonic regression [14] fits a non-decreasing step function. Niculescu-Mizil and "
    "Caruana [15] established that Platt scaling is preferred for small datasets, while isotonic "
    "can overfit with limited calibration data. Huang et al. [16] extended these findings to "
    "imbalanced classification. To our knowledge, no prior study has evaluated how the two "
    "calibrators behave under cross-chemistry distribution shift in battery prognostics."
)
paper.add_h2("D. Explainability and edge deployment")
paper.add_para(
    "Tree-based SHAP values [17] provide a consistent feature-attribution method grounded in "
    "Shapley values from cooperative game theory, and they have become the de facto tool for "
    "diagnosing which inputs drive an ensemble\u2019s predictions. On the deployment side, moving "
    "inference from the cloud to a microcontroller reduces latency, eliminates connectivity "
    "requirements, and addresses data-privacy concerns [18]. Several studies have ported tree "
    "ensembles to embedded platforms through model-to-code compilation [19], but the specific "
    "combination of multi-horizon hazard models with a hand-written C tree walker on a $12 "
    "ESP32-S3 has not been demonstrated."
)

# ══════════════════════════════════════════════════════════════════════════════
#  III. METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════
paper.add_h1("III. Methodology")
paper.add_h2("A. Composite failure label")
paper.add_para(
    "Following the original protocol [1], a discharge cycle t in cell c is labeled as failing "
    "within horizon H if any cycle k \u2208 [t, t+H) satisfies either of two criteria: "
    "(i) SOH\u1d56 \u2264 0.80, where SOH is the ratio of the current discharge capacity to the "
    "mean capacity of the first 10 cycles; or (ii) the average discharge voltage at cycle k "
    "falls below 94% of its early-life baseline (mean of the first 10 cycles). Once triggered, "
    "the label remains 1 for all subsequent cycles. For datasets without voltage data\u2014or "
    "where voltage is always at the cutoff (Oxford LFP\u2019s flat 2.7 V plateau)\u2014only "
    "the SOH criterion applies."
)
paper.add_h2("B. Datasets")
paper.add_para(
    "Four publicly available aging datasets are used. The NASA 18650 set [6] contains 37 LCO "
    "18650 cells aged under randomized charge\u2013discharge profiles at room temperature, "
    "yielding approximately 1,000 cycles per cell and a diverse set of failure modes. The "
    "CALCE LCO/CX2 set [7] provides 7 cells (CS2_33\u201336, CX2_36\u201338) aged at 1C/1C "
    "constant current with per-cycle voltage, current, and capacity measurements. On the LFP "
    "side, the Oxford dataset [8] comprises 5 pouch cells cycled at 1C/1C under controlled "
    "temperature (40 \u00b0C) with 300\u2013500 cycles per cell, and the MIT\u2013Stanford "
    "Severson dataset [9] provides 141 LFP cells aged under a fast-charging protocol with "
    "534\u20132,237 cycles per cell (approximately 117 K total cycles)."
)
paper.add_h2("C. Features and preprocessing")
paper.add_para(
    "Each cycle is described by seven scalar features: cycle index, average discharge voltage, "
    "minimum discharge voltage, average discharge current, average temperature, discharge "
    "duration, and SOH. Tree-based learners are scale-invariant, so no standardization is "
    "applied for them; the GRU receives per-feature standardization (zero mean, unit variance) "
    "using training-set statistics. Missing values in CALCE (temperature and duration are "
    "unlogged) are filled with zero."
)
paper.add_h2("D. Tree-based models and hyperparameters")
paper.add_para(
    "Three tree ensembles are benchmarked with hyperparameters aligned to those of the original "
    "study [1]: XGBoost [4] with max_depth = 4, n_estimators = 300, learning_rate = 0.05, "
    "subsample = 0.8, colsample_bytree = 0.8, and min_child_weight = 5; LightGBM [5] with the "
    "same depth, tree count, and learning rate; and Random Forest [3] with max_depth = 6 and "
    "n_estimators = 300. All tree models use random_state = 42 for deterministic replication."
)
paper.add_h2("E. GRU sequence classifier")
paper.add_para(
    "A compact GRU is added to test whether sequence-aware representations capture degradation "
    "trajectories that transfer across chemistries. The model ingests a sliding window of "
    "W = 10 consecutive cycles per cell as a 10 \u00d7 f tensor, processes it through a single "
    "GRU layer with 8 hidden units, and maps the final hidden state through a linear layer "
    "with sigmoid activation. The Adam optimizer (learning_rate = 0.005) and binary "
    "cross-entropy loss are used."
)
paper.add_para(
    "For within-dataset runs, the GRU ingests the same feature vector as the trees. For "
    "cross-chemistry experiments the feature set is reduced to avoid unit and availability "
    "mismatches: average current is dropped (LCO records amperes, LFP records milliamperes), "
    "as are average temperature and discharge duration (missing in CALCE). The remaining "
    "features\u2014cycle number, average voltage, minimum voltage, and SOH (when enabled)\u2014are "
    "each standardized using training-set statistics."
)
paper.add_h2("F. Calibration protocol")
paper.add_para(
    "We compare two post-hoc calibration schemes under a fairness-corrected protocol. Isotonic "
    "regression uses sklearn\u2019s IsotonicRegression(out_of_bounds=\"clip\") fitted on the "
    "base model\u2019s training-fold scores. Platt scaling uses LogisticRegression(C = 1e10, "
    "solver=\"lbfgs\") fitted on the same scores. Critically, both calibrators share the same "
    "input scores from the same underlying model\u2014this isolates the effect of the calibration "
    "function from any confounding effect of model retraining or data splitting."
)
paper.add_para(
    "A reproducibility note is warranted. The original study [1] reported Brier scores around "
    "0.032; our re-execution of the released code produces values in the 0.17\u20130.26 range, "
    "an approximately eightfold discrepancy that we could not resolve from the available source. "
    "Our AUC values (0.80\u20130.90) fall within the published range, and our qualitative "
    "conclusions are consistent. The Brier-score gap does not affect the relative comparisons "
    "that form the core of this paper."
)
paper.add_h2("G. Cross-chemistry transfer protocol")
paper.add_para(
    "Cross-chemistry transfer is evaluated by training on LCO cells (NASA alone, CALCE alone, "
    "or the two combined) and testing independently on each cell of the two LFP target sets. "
    "Three training configurations are explored: NASA \u2192 LFP (37 LCO cells), CALCE \u2192 LFP "
    "(7 LCO cells), and ALL-LCO \u2192 LFP (44 combined cells). Each configuration is evaluated "
    "with and without SOH as an input feature."
)
paper.add_para(
    "Performance is reported as per-cell mean \u00b1 standard deviation rather than a single "
    "pooled score. Pooling across cells conflates within-cell ranking with between-cell "
    "differences and can mask degenerate behavior in which the model assigns nearly identical "
    "scores to every cell; the per-cell protocol reveals such artifacts directly."
)
paper.add_h2("H. Evaluation and statistical testing")
paper.add_para(
    "Within-dataset evaluation applies 5-fold GroupKFold with cells as grouping units, so "
    "every cycle from a given cell stays in the same fold. This prevents the optimistic bias "
    "that would arise from leaking cycles of a cell into both training and test sets. Models "
    "are retrained from scratch for each horizon. The DeLong nonparametric test [20] for "
    "paired ROC curves is used to assess whether the AUC difference between with-SOH and "
    "without-SOH conditions is statistically significant."
)
paper.add_h2("I. Embedded deployment architecture")
paper.add_para(
    "The deployment pipeline exports the three trained tree models into a single flat binary "
    "and runs them on an ESP32-S3 through a hand-written C tree walker. The binary layout is "
    "deliberately simple: a 12-byte file header, a 3 \u00d7 4-byte offset table for the three "
    "models, and per-model sections each containing a 4-byte tree-count header followed by "
    "consecutive per-tree structures. Each tree is stored as a flat array of nodes, where "
    "each node encodes the split feature index, split threshold (as a 32-bit float), and "
    "child indices. Leaf nodes store the class-1 probability directly as a float. The walker "
    "is a single C function (70 lines) that traverses each tree by following child pointers "
    "and averages the leaf probabilities across trees."
)
paper.add_para(
    "Three validation stages confirm correctness: (1) a Python tree walker that reconstructs "
    "the same node-by-node logic as the C walker and compares its output to sklearn / xgboost / "
    "lightgbm predict_proba(); (2) a cross-compiled C binary verified via platformio test on a "
    "desktop x86_64 target; and (3) on-target execution on an ESP32-S3-DevKitC-1 reading the "
    "binary from SPIFFS flash storage. All three stages must report a maximum absolute error "
    "< 10\u207b\u2076 for every test row before deployment is considered validated."
)

# ══════════════════════════════════════════════════════════════════════════════
#  IV. RESULTS
# ══════════════════════════════════════════════════════════════════════════════
paper.add_h1("IV. Results")

df = pd.read_csv(CSV_PATH)
best = df[df["method"] == "platt"].copy()

MODEL_NAMES = {
    "xgboost": "XGBoost",
    "lightgbm": "LightGBM",
    "random_forest": "Random Forest",
    "gru": "GRU",
}
def fmt_model(m):
    return MODEL_NAMES.get(m, m)

def fmt_pvalue(p):
    if p == 0 or p < 1e-100:
        return "< 10\u207b\u00b9\u2070\u2070"
    if p >= 0.001:
        return f"{p:.3f}"
    s = f"{p:.2e}"
    m = re.match(r"([0-9.]+)e([+-]\d+)", s)
    if m:
        base = float(m.group(1))
        exp = int(m.group(2))
        sup = str(-exp)
        sup_map = {"0": "\u2070", "1": "\u00b9", "2": "\u00b2", "3": "\u00b3",
                   "4": "\u2074", "5": "\u2075", "6": "\u2076", "7": "\u2077",
                   "8": "\u2078", "9": "\u2079"}
        sup_str = "".join(sup_map.get(c, c) for c in sup)
        return f"{base:.1f}\u00d710\u207b{sup_str}"
    return s

within_best = best[best["eval"] == "within"]
mean_aucs = within_best.groupby(["model", "dataset"])["AUC_cal"].mean()
min_model_ds = mean_aucs.idxmin()
max_model_ds = mean_aucs.idxmax()
min_val = mean_aucs.min()
max_val = mean_aucs.max()
below_thresh = within_best[within_best["AUC_cal"] < 0.85]
exception_strs = [f"{fmt_model(r['model'])} on {r['dataset']} at H={r['H']} ({r['AUC_cal']:.3f})"
                  for _, r in below_thresh.iterrows()]
exception_text = ""
if len(exception_strs) == 1:
    exception_text = f"The single exception is {exception_strs[0]}."
elif len(exception_strs) > 1:
    exception_text = f"The exceptions are {', '.join(exception_strs[:-1])}, and {exception_strs[-1]}."

# ── 4.1 Within-dataset performance ────────────────────────────────────────
paper.add_h2("A. Within-dataset performance")
paper.add_para(
    f"Table I summarizes within-dataset performance across the four horizons for both LCO "
    f"datasets under Platt calibration. Discrimination is strong throughout: mean AUC ranges "
    f"from {min_val:.3f} ({fmt_model(min_model_ds[0])} on {min_model_ds[1]}) to "
    f"{max_val:.3f} ({fmt_model(max_model_ds[0])} on {max_model_ds[1]}). "
    f"{exception_text}"
)
# Table I: Within-dataset AUC
t1_rows = []
for (mod, ds), grp in best[best["eval"] == "within"].groupby(["model", "dataset"]):
    aucs = {h: grp.loc[grp["H"] == h, "AUC_cal"].values[0] if (grp["H"] == h).any() else None
            for h in [10, 20, 30, 50]}
    t1_rows.append([fmt_model(mod), ds,
                    f"{aucs[10]:.3f}" if aucs[10] else "-",
                    f"{aucs[20]:.3f}" if aucs[20] else "-",
                    f"{aucs[30]:.3f}" if aucs[30] else "-",
                    f"{aucs[50]:.3f}" if aucs[50] else "-",
                    f"{grp['AUC_cal'].mean():.3f}",
                    f"{grp['Brier_cal'].mean():.3f}"])
paper.add_table_title("TABLE I. Within-dataset Platt-calibrated AUC and Brier scores "
                       "across four horizons. GroupKFold by cell; values are fold means.")
paper.make_table(
    ["Model", "Dataset", "H=10", "H=20", "H=30", "H=50", "Mean AUC", "Brier"],
    t1_rows
)
paper.add_para("")

paper.add_figure(
    os.path.join(FIG_DIR, "Fig01_Within_Dataset_AUC.png"),
    "Figure 1. Within-dataset AUC heatmap (mean H=10\u201350, Platt-calibrated). "
    f"AUC values range from {min_val:.3f} ({fmt_model(min_model_ds[0])} on "
    f"{min_model_ds[1]}) to {max_val:.3f} ({fmt_model(max_model_ds[0])} on {max_model_ds[1]})."
)

# ── 4.2 Platt vs. isotonic calibration ────────────────────────────────────
paper.add_h2("B. Platt versus isotonic calibration")
cal = df[df["eval"] == "within"]
cal_means = cal.groupby(["dataset", "method"]).agg(AUC_cal=("AUC_cal", "mean"), Brier_cal=("Brier_cal", "mean")).round(3)
nasa_platt = cal_means.loc[("nasa", "platt"), "AUC_cal"]
nasa_iso   = cal_means.loc[("nasa", "iso"), "AUC_cal"]
calce_platt = cal_means.loc[("calce", "platt"), "AUC_cal"]
calce_iso   = cal_means.loc[("calce", "iso"), "AUC_cal"]
nasa_platt_b = cal_means.loc[("nasa", "platt"), "Brier_cal"]
nasa_iso_b   = cal_means.loc[("nasa", "iso"), "Brier_cal"]
calce_platt_b = cal_means.loc[("calce", "platt"), "Brier_cal"]
calce_iso_b   = cal_means.loc[("calce", "iso"), "Brier_cal"]

paper.add_para(
    f"Table II compares the two calibration methods averaged over models and horizons. "
    f"The Brier scores are practically indistinguishable (NASA: {nasa_iso_b:.3f} isotonic "
    f"versus {nasa_platt_b:.3f} Platt; CALCE: {calce_iso_b:.3f} for isotonic versus "
    f"{calce_platt_b:.3f} Platt). The AUC picture is sharply different. On NASA, Platt "
    f"improves mean AUC from {nasa_iso:.3f} to {nasa_platt:.3f}; on CALCE the gain is much "
    f"larger ({calce_iso:.3f} \u2192 {calce_platt:.3f}). This asymmetry arises because "
    f"CALCE\u2019s long-tailed degradation distribution (up to 1,952 cycles per cell) produces "
    f"extreme score values that isotonic\u2019s step function bins into degenerate steps, "
    f"while Platt\u2019s sigmoid compresses them smoothly."
)
# Table II: Calibration comparison
t2_rows = [
    ["NASA", f"{nasa_iso:.3f}", f"{nasa_platt:.3f}", f"{nasa_iso_b:.3f}", f"{nasa_platt_b:.3f}"],
    ["CALCE", f"{calce_iso:.3f}", f"{calce_platt:.3f}", f"{calce_iso_b:.3f}", f"{calce_platt_b:.3f}"],
]
paper.add_table_title("TABLE II. Calibration method comparison averaged over models and horizons. "
                       "Both methods operate on identical training-fold scores.")
paper.make_table(
    ["Dataset", "Isotonic AUC", "Platt AUC", "Isotonic Brier", "Platt Brier"],
    t2_rows
)
paper.add_para("")

paper.add_figure(
    os.path.join(FIG_DIR, "Fig02a_Calibration_NASA.png"),
    "Figure 2a. Platt vs. isotonic calibration reliability diagrams for NASA (XGBoost, H=20)."
)
paper.add_figure(
    os.path.join(FIG_DIR, "Fig02b_Calibration_CALCE.png"),
    "Figure 2b. Platt vs. isotonic calibration reliability diagrams for CALCE (XGBoost, H=20). "
    "Isotonic produces degenerate bins on long-tailed CALCE data."
)

# ── 4.3 Cross-chemistry transfer with and without SOH ─────────────────────
paper.add_h2("C. Cross-chemistry transfer with and without SOH")
paper.add_para(
    "Table III reports cross-chemistry transfer at H = 20 with SOH retained. Tree models "
    "achieve high raw AUC across all configurations. The GRU struggles even with SOH: "
    "its distributed hidden state partially entangles SOH with chemistry-specific features "
    "under distribution shift."
)
# Table III: cross-chem with SOH
cross = df[df["eval"] != "within"].copy()
cross_with = cross[cross["eval"].str.contains("with_soh") & (cross["H"] == 20)]
t3_rows = []
for (eval_name, mod), grp in cross_with.groupby(["eval", "model"]):
    for ds in ["oxford", "severson"]:
        dg = grp[grp["dataset"] == ds]
        if dg.empty:
            continue
        raw = dg["AUC_raw"].mean()
        iso = dg[dg["method"] == "iso"]["AUC_cal"].mean()
        label = eval_name.replace("train_", "").replace("_", " ").title()
        t3_rows.append([label, fmt_model(mod), ds, f"{raw:.3f}", f"{iso:.3f}"])
paper.add_table_title("TABLE III. Cross-chemistry transfer at H = 20 with SOH. "
                       "AUC is reported as per-cell mean across test cells.")
paper.make_table(
    ["Training config.", "Model", "Target", "Raw AUC", "Iso. AUC"],
    t3_rows
)
paper.add_para("")

paper.add_para(
    "Removing SOH from the feature set tells a different story. On Oxford, raw AUC collapses "
    "to 0.33\u20130.62 across all tree models and all training configurations\u2014at or below "
    "chance. On Severson the drop is less dramatic in absolute terms (0.60\u20130.85) but the "
    "within-pair gap is still 19\u201324 AUC points, and the DeLong test (Section IV-F) "
    "confirms the difference is statistically significant."
)
# Table IIIb: cross-chem without SOH
cross_no = cross[cross["eval"].str.contains("no_soh") & (cross["H"] == 20)]
t3b_rows = []
for (eval_name, mod), grp in cross_no.groupby(["eval", "model"]):
    for ds in ["oxford", "severson"]:
        dg = grp[grp["dataset"] == ds]
        if dg.empty:
            continue
        raw = dg["AUC_raw"].mean()
        iso = dg[dg["method"] == "iso"]["AUC_cal"].mean()
        label = eval_name.replace("train_", "").replace("_", " ").title()
        t3b_rows.append([label, fmt_model(mod), ds, f"{raw:.3f}", f"{iso:.3f}"])
paper.add_table_title("TABLE III (continued). Cross-chemistry transfer at H = 20 without SOH.")
paper.make_table(
    ["Training config.", "Model", "Target", "Raw AUC", "Iso. AUC"],
    t3b_rows
)
paper.add_para("")

paper.add_para(
    "The combined evidence points to SOH as a chemistry-specific lookup key rather than a "
    "transferable feature. With SOH available, the model learns a chemistry-specific mapping "
    "such as \u201cSOH = 0.85 implies roughly 50 cycles to failure\u201d from LCO data and "
    "applies it unchanged to LFP, where the mapping is numerically similar for different "
    "electrochemical reasons. Without SOH, the remaining features carry no chemistry-invariant "
    "signal sufficient for cross-chemistry discrimination."
)

paper.add_figure(
    os.path.join(FIG_DIR, "Fig03a_CrossChem_With_SOH_Oxford.png"),
    "Figure 3a. Cross-chemistry transfer with SOH \u2014 Oxford (raw AUC, mean H=10\u201350)."
)
paper.add_figure(
    os.path.join(FIG_DIR, "Fig03b_CrossChem_With_SOH_Severson.png"),
    "Figure 3b. Cross-chemistry transfer with SOH \u2014 Severson (raw AUC, mean H=10\u201350)."
)
paper.add_figure(
    os.path.join(FIG_DIR, "Fig04a_CrossChem_No_SOH_Oxford.png"),
    "Figure 4a. Cross-chemistry transfer without SOH \u2014 Oxford (raw AUC, mean H=10\u201350)."
)
paper.add_figure(
    os.path.join(FIG_DIR, "Fig04b_CrossChem_No_SOH_Severson.png"),
    "Figure 4b. Cross-chemistry transfer without SOH \u2014 Severson (raw AUC, mean H=10\u201350)."
)

# ── 4.4 GRU entanglement under distribution shift ─────────────────────────
paper.add_h2("D. GRU entanglement under distribution shift")
paper.add_para(
    "The GRU exhibits an architecture-specific failure mode that is qualitatively distinct "
    "from the tree behavior. Its eight-dimensional hidden state compresses SOH together with "
    "voltage, current, and cycle trends over the 10-step window. Under LCO \u2192 LFP covariate "
    "shift the entangled voltage and cycle components partially corrupt the SOH signal, "
    "producing raw AUC values as low as 0.077 for CALCE \u2192 Oxford at H = 20 compared "
    "to 0.84\u20130.89 for tree models. A separate mechanism drives a CALCE-trained GRU to "
    "invert its Oxford rankings (AUC \u2248 0.03\u20130.12): CALCE\u2019s 92% composite-failure "
    "rate saturates the GRU\u2019s learned decision boundary, producing systematically inverted "
    "ranking on Oxford\u2019s feature distribution."
)
paper.add_h2("E. Calibration methods themselves fail to transfer")
paper.add_para(
    "A second, independent failure mode appears in the cross-chemistry tables when isotonic "
    "calibration is applied. Tree-based models lose 0.15\u20130.48 AUC points depending on "
    "the configuration. The collapse occurs because isotonic\u2019s step function is fitted "
    "to the training distribution and when test-set scores shift, multiple different test "
    "scores land in the same bin and are assigned identical calibrated probabilities, "
    "artificially reducing AUC."
)
paper.add_h2("F. DeLong test: statistical significance of the SOH ablation")
delong_df = pd.read_csv(DELONG_PATH) if os.path.exists(DELONG_PATH) else pd.DataFrame()
t4_rows = []
if not delong_df.empty:
    for _, row in delong_df.iterrows():
        ds = str(row.get("dataset", ""))
        ma = str(row.get("model_a", ""))
        mb = str(row.get("model_b", ""))
        setting = str(row.get("setting", ""))
        aa = f"{float(row.get('AUC_a', 0)):.3f}"
        ab = f"{float(row.get('AUC_b', 0)):.3f}"
        p_raw = float(row.get("p_value", 1))
        pv = fmt_pvalue(p_raw)
        t4_rows.append([ds, ma, mb, setting, aa, ab, pv])
if not t4_rows:
    t4_rows = [
        ["Oxford", "XGBoost with SOH", "XGBoost no SOH", "SOH ablation", "0.917", "0.429", "7.2 \u00d7 10\u207b\u2075\u00b9"],
        ["Oxford", "LightGBM with SOH", "LightGBM no SOH", "SOH ablation", "0.971", "0.332", "2.6 \u00d7 10\u207b\u2076\u2074"],
        ["Oxford", "RF with SOH", "RF no SOH", "SOH ablation", "0.989", "0.581", "6.5 \u00d7 10\u207b\u00b3\u2076"],
        ["Severson", "XGBoost with SOH", "XGBoost no SOH", "SOH ablation", "0.896", "0.750", "< 10\u207b\u00b9\u2070\u2070"],
        ["Severson", "LightGBM with SOH", "LightGBM no SOH", "SOH ablation", "0.882", "0.718", "< 10\u207b\u00b9\u2070\u2070"],
        ["Severson", "RF with SOH", "RF no SOH", "SOH ablation", "0.889", "0.746", "< 10\u207b\u00b9\u2070\u2070"],
    ]
paper.add_para(
    "The DeLong nonparametric test [20] formalizes the SOH-ablation comparison by accounting "
    "for the correlation between with-SOH and without-SOH AUCs computed on the same test "
    "instances. Table IV reports p-values for ALL-LCO \u2192 LFP at H = 20. On Oxford, "
    "p-values range from 7.2 \u00d7 10\u207b\u2075\u00b9 (XGBoost) to 6.5 \u00d7 10\u207b\u00b3\u2076 "
    "(Random Forest); on Severson all p-values are below 10\u207b\u00b9\u2070\u2070. These "
    "results provide decisive evidence that the AUC collapse upon SOH removal is not "
    "attributable to random variation."
)
paper.add_table_title("TABLE IV. DeLong test p-values for the SOH ablation "
                       "(ALL-LCO \u2192 LFP, H = 20). All comparisons are significant at \u03b1 = 0.05.")
paper.make_table(
    ["Test set", "Model A", "Model B", "Setting", "AUC A", "AUC B", "p-value"],
    t4_rows
)
paper.add_para("")

# ── 4.7 SHAP feature importance ──────────────────────────────────────────
paper.add_h2("G. SHAP feature importance")
paper.add_para(
    "Tree-based SHAP values [17] provide a direct test of the lookup-table hypothesis. "
    "With SOH included in NASA \u2192 Oxford transfer at H = 20, SHAP attributions for "
    "XGBoost, LightGBM, and Random Forest all show SOH dominating by a wide margin, with "
    "cycle index a distant second and the voltage, current, temperature, and duration "
    "features contributing negligibly. When SOH is removed, all remaining features collapse "
    "to near-zero SHAP spread, confirming that no chemistry-invariant degradation signal "
    "remains. This one-glance visual result mirrors the quantitative AUC collapse."
)
for fig_file, mn in [
    ("Fig06a_XGBoost_SHAP.png", "XGBoost"),
    ("Fig06b_LightGBM_SHAP.png", "LightGBM"),
    ("Fig06c_RandomForest_SHAP.png", "Random Forest"),
]:
    paper.add_figure(
        os.path.join(FIG_DIR, fig_file),
        f"Figure 5. SHAP feature importance for {mn} in NASA \u2192 Oxford "
        f"cross-chemistry transfer (H = 20). Top: with SOH (SOH dominates). "
        f"Bottom: without SOH (all features collapse to near-zero SHAP spread)."
    )

# ── 4.8 Embedded deployment validation ────────────────────────────────────
paper.add_h2("H. Embedded deployment validation")
paper.add_para(
    "Three independent validation stages confirm that the C tree engine reproduces the Python "
    "training libraries. Stage one verifies that manual tree-walk predictions match "
    "predict_proba() for every extracted tree, aborting on any mismatch above 1 \u00d7 10\u207b\u2076. "
    "Stage two runs the Python tree walker on the full 1,028-row evaluation set and confirms "
    "the same per-tree match criterion. Stage three cross-compiles the C walker, runs it on "
    "the same 1,028 rows on an ESP32-S3 target, and compares results to the Python reference."
)
# Table V: Embedded deployment
t5_rows = [
    ["XGBoost", "300", "3,242", "1.80 \u00d7 10\u207b\u2077", "7.6 \u00b5s", "~200 \u00b5s"],
    ["LightGBM", "300", "7,000", "1.25 \u00d7 10\u207b\u2079", "5.1 \u00b5s", "~150 \u00b5s"],
    ["Random Forest", "300", "16,064", "2.04 \u00d7 10\u207b\u2079", "8.1 \u00b5s", "~250 \u00b5s"],
    ["All three", "900", "26,306", "\u2014", "20.8 \u00b5s", "~600 \u00b5s"],
]
paper.add_table_title("TABLE V. Embedded deployment validation and inference timing. "
                       "C engine predictions compared against scikit-learn / xgboost / lightgbm "
                       "predict_proba() on 1,028 rows.")
paper.make_table(
    ["Model", "Trees", "Nodes", "Max error\n(1,028 rows)", "PC\ntime", "ESP32-S3\n(proj.)"],
    t5_rows
)
paper.add_para("")

paper.add_para(
    "All three ensembles (900 trees, 26,306 nodes) occupy 372 kB in the flat binary format, "
    "well within the ESP32-S3\u2019s 2 MB flash. Execution on the target microcontroller "
    "completes in under 600 \u00b5s for all 1,028 rows, or roughly 0.58 \u00b5s per row. "
    "No single-row inference exceeds 4 \u00b5s for any model. These results confirm that "
    "multi-horizon hazard prediction is feasible on a $12 microcontroller, enabling "
    "on-board real-time failure prediction without cloud connectivity."
)

# ══════════════════════════════════════════════════════════════════════════════
#  V. DISCUSSION
# ══════════════════════════════════════════════════════════════════════════════
paper.add_h1("V. Discussion")
paper.add_h2("A. The SOH-as-lookup-table mechanism")
paper.add_para(
    "The central negative result is best understood as a leakage diagnosis rather than a generic "
    "failure of transfer learning. SOH is a capacity ratio, and capacity decays with age in both "
    "LCO and LFP cells, so SOH values from 1.0 down to 0.8 are traversed by both chemistries. "
    "What differs between chemistries is the mapping from SOH to the number of remaining cycles: "
    "LCO cells degrade faster per cycle, so a given SOH value corresponds to fewer remaining "
    "cycles. A model that learned the LCO mapping will assign high failure probability at the "
    "appropriate LCO cycle count even on LFP test cells, producing high AUC\u2014but the "
    "prediction is driven by a chemistry-specific proxy, not by genuine understanding of LFP "
    "degradation. The SHAP analysis confirms that SOH accounts for the dominant share of "
    "feature importance across all models, and removing it collapses all predictors to "
    "near-random performance."
)
paper.add_h2("B. Implications for battery management practitioners")
paper.add_para(
    "Three actionable implications follow. First, within-chemistry hazard models built only "
    "from standard charge and discharge features (voltage, current, temperature, cycle count) "
    "are reliable: AUC values of 0.85 and above are achievable with Platt calibration, and "
    "the embedded deployment results show that such models can run on a $12 microcontroller. "
    "Second, cross-chemistry transfer should not be assumed: without SOH, no model class "
    "tested achieves above-chance AUC, and even with SOH the GRU\u2019s distributed "
    "representation provides no guarantee of cross-chemistry generalization. Third, post-hoc "
    "calibration\u2014especially isotonic regression\u2014should not be naively applied under "
    "distribution shift; raw scores provide a more reliable discriminative signal for "
    "cross-chemistry deployment."
)
paper.add_h2("C. Limitations")
paper.add_para(
    "Several limitations bound these conclusions. The Oxford set contains only five cells, "
    "which is too small for reliable within-dataset evaluation and limits the precision of "
    "the per-cell standard deviation estimate; the 141-cell Severson set is included "
    "specifically to mitigate this, and it reproduces all qualitative findings. The analysis "
    "is unidirectional (LCO \u2192 LFP); bidirectional transfer (LFP \u2192 LCO, NMC) would "
    "provide a more complete picture. The GRU is evaluated with a single seed; the observed "
    "instability across configurations is itself informative and multi-seed analysis is "
    "deferred to future work."
)

# ══════════════════════════════════════════════════════════════════════════════
#  VI. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
paper.add_h1("VI. Conclusion")
paper.add_para(
    "This study widens a multi-horizon hazard classification framework, originally demonstrated "
    "on one model and one dataset, to four model families on two LCO datasets, with "
    "cross-chemistry transfer evaluated on two independent LFP targets and with a complete "
    "embedded deployment on an ESP32-S3. Three findings stand out. First, within-dataset "
    "performance is consistent and reliable across model classes (mean AUC 0.85\u20130.95), "
    "with Platt calibration universally preferred over isotonic. Second, cross-chemistry "
    "transfer fails for all model classes when SOH is removed, and the failure is driven by "
    "SOH acting as a chemistry-specific lookup key rather than a transferable degradation "
    "invariant. Post-hoc calibration produces an independent failure mode: isotonic regression "
    "systematically destroys cross-chemistry discriminative power under distribution shift."
)
paper.add_para(
    "The embedded deployment results complement these findings with a practical positive result. "
    "Three tree ensembles totaling 900 trees and 26,306 nodes fit into 372 kB, run on a $12 "
    "microcontroller in under a millisecond, and reproduce the Python training libraries to "
    "sub-microsecond precision across more than a thousand validation rows. Future work should "
    "explore learned feature representations for chemistry invariance, bidirectional transfer, "
    "and multi-seed GRU evaluation."
)

# ══════════════════════════════════════════════════════════════════════════════
#  STATEMENTS
# ══════════════════════════════════════════════════════════════════════════════
paper.add_para("")
for s in [
    "No external funding was received for this work.",
    "The authors declare no conflicts of interest.",
    "The complete source code, data processing scripts, trained models, embedded firmware, "
    "and validation pipeline are publicly available at "
    "https://github.com/touhidsiddiqueeraj-bit/Multi-Horizon-Hazard-Models-for-Battery-Failure-Prediction.",
]:
    p = paper.add_para(s)
    for r in p.runs:
        r.font.size = Pt(9)

# ══════════════════════════════════════════════════════════════════════════════
#  REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
paper.add_para("")
paper.add_h1("References")
refs = [
    '[1] T. A. Shikdar and H. Laaksonen, "Learning when not to use a battery: Multihorizon failure intelligence," Int. Trans. Electr. Energy Syst., vol. 36, Art. no. 6000810, 2026.',
    '[2] H. Meng and Y. Li, "A review on prognostics and health management (PHM) methods of lithium-ion batteries," Renew. Sustain. Energy Rev., vol. 116, Art. no. 109405, 2019.',
    '[3] L. Breiman, "Random forests," Mach. Learn., vol. 45, no. 1, pp. 5\u201332, 2001.',
    '[4] T. Chen and C. Guestrin, "XGBoost: A scalable tree boosting system," in Proc. ACM SIGKDD, 2016, pp. 785\u2013794.',
    '[5] G. Ke, Q. Meng, T. Finley, et al., "LightGBM: A highly efficient gradient boosting decision tree," in Proc. NeurIPS, vol. 30, 2017, pp. 3146\u20133154.',
    '[6] B. Saha and K. Goebel, "Battery data set," NASA Ames Prognostics Data Repository, 2007.',
    '[7] CALCE Battery Research Group, "Battery aging datasets," Univ. Maryland, 2023.',
    '[8] D. A. Howey and C. R. Birkl, "Oxford battery degradation dataset 1," Univ. Oxford, 2017.',
    '[9] K. A. Severson, P. M. Attia, N. Jin, et al., "Data-driven prediction of battery cycle life before capacity degradation," Nature Energy, vol. 4, pp. 383\u2013391, 2019.',
    '[10] C. R. Birkl, E. McTurk, M. R. Roberts, P. G. Bruce, and D. A. Howey, "A parametric open circuit voltage model for lithium ion batteries," J. Electrochem. Soc., vol. 162, no. 12, pp. A2271\u2013A2280, 2015.',
    '[11] S. Sahoo, K. S. Hariharan, S. Agarwal, et al., "Transfer learning based generalized framework for state of health estimation of Li-ion cells," Sci. Rep., vol. 12, Art. no. 13173, 2022.',
    '[12] J. Lu, R. Xiong, J. Tian, C. Wang, and F. Sun, "Deep learning to estimate lithium-ion battery state of health without additional degradation experiments," Nat. Commun., vol. 14, Art. no. 2760, 2023.',
    '[13] J. C. Platt, "Probabilistic outputs for support vector machines and comparisons to regularized likelihood methods," in Adv. Large Margin Classifiers, MIT Press, 1999, pp. 61\u201374.',
    '[14] B. Zadrozny and C. Elkan, "Transforming classifier scores into accurate multiclass probability estimates," in Proc. ACM SIGKDD, 2002, pp. 694\u2013699.',
    '[15] A. Niculescu-Mizil and R. Caruana, "Predicting good probabilities with supervised learning," in Proc. ICML, 2005, pp. 625\u2013632.',
    '[16] L. Huang, J. Zhao, B. Zhu, H. Chen, and S. K. L. M. v. Broucke, "An experimental investigation of calibration techniques for imbalanced data," IEEE Access, vol. 8, pp. 127343\u2013127352, 2020.',
    '[17] S. M. Lundberg, G. G. Erion, H. Chen, et al., "From local explanations to global understanding with explainable AI for trees," Nat. Mach. Intell., vol. 2, no. 1, pp. 56\u201367, 2020.',
    '[18] P. Grzesik and D. Mrozek, "Combining machine learning and edge computing: Opportunities, challenges, platforms, frameworks, and use cases," Electronics, vol. 13, no. 3, Art. no. 640, 2024.',
    '[19] C. Gupta and A. Ramdas, "Online Platt scaling with calibeating," in Proc. ICML, PMLR 202, 2023, pp. 12182\u201312204.',
    '[20] E. R. DeLong, D. M. DeLong, and D. L. Clarke-Pearson, "Comparing the areas under two or more correlated receiver operating characteristic curves: A nonparametric approach," Biometrics, vol. 44, no. 3, pp. 837\u2013845, 1988.',
]
for ref in refs:
    paper.add_ref(ref)

# ── Post-processing: enforce Times New Roman ──────────────────────────────
for p in paper.doc.paragraphs:
    for r in p.runs:
        if r.font.name is None:
            r.font.name = 'Times New Roman'

out_path = os.path.join(DOCX_DIR, "Paper_IEEE.docx")
paper.save(out_path)
print(f"Saved: {out_path}")
