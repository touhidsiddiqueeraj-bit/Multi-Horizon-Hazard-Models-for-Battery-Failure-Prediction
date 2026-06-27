"""IEEE Access document builder — two-column layout, running headers, proper styles."""
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def make_style(doc, name, base='Normal', size=10, bold=False, italic=False,
               alignment=None, font_name='Times New Roman', sb=0, sa=3, caps=False):
    st = doc.styles.add_style(name, 1)
    try:
        st.base_style = doc.styles[base]
    except KeyError:
        pass
    st.font.name = font_name
    st.font.size = Pt(size)
    st.font.bold = bold
    st.font.italic = italic
    if caps:
        rPr = st.element.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
            st.element.append(rPr)
        rPr.append(parse_xml(f'<w:smallCaps {nsdecls("w")}/>'))
    if alignment is not None:
        st.paragraph_format.alignment = alignment
    st.paragraph_format.space_before = Pt(sb)
    st.paragraph_format.space_after = Pt(sa)

def add_page_number(paragraph):
    run = paragraph.add_run()
    run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    run2 = paragraph.add_run()
    run2._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    run3 = paragraph.add_run()
    run3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

class IEEEPaper:
    def __init__(self):
        self.doc = Document()
        self._setup_page()
        self._create_styles()
        self._title_section = self.doc.sections[0]
        self._body_section = None

    def _setup_page(self):
        s = self.doc.sections[0]
        s.page_width = Inches(8.5)
        s.page_height = Inches(11)
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.69)
        s.right_margin = Inches(0.69)
        s.different_first_page_header_footer = True

        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(3)

    def _create_styles(self):
        make_style(self.doc, 'PaperTitle', size=24, bold=True,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=6)
        make_style(self.doc, 'AU', size=10,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=3)
        make_style(self.doc, 'PI', size=10,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=3)
        make_style(self.doc, 'H1', size=10, bold=True, caps=True,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, sb=12, sa=3)
        make_style(self.doc, 'H2', size=10, bold=True,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT, sb=10, sa=3)
        make_style(self.doc, 'H3', size=10, italic=True,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=3)
        make_style(self.doc, 'FigCaption', size=9, bold=True,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, sb=4, sa=6)
        make_style(self.doc, 'TableTitle', size=9, bold=True,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, sb=6, sa=3)
        make_style(self.doc, 'RefStyle', size=9,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT, sa=2)
        make_style(self.doc, 'IndexTerms', size=9,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT, sa=6)

    def begin_body(self, running_title="", volume="XX"):
        """Insert section break and start two-column body."""
        # Title page (section 0) should have EMPTY first-page header/footer
        fph = self._title_section.first_page_header
        fph.paragraphs[0].text = ""
        fpf = self._title_section.first_page_footer
        fpf.paragraphs[0].text = ""

        ns = self.doc.add_section(WD_SECTION_START.CONTINUOUS)
        ns.top_margin = Inches(0.75)
        ns.bottom_margin = Inches(0.75)
        ns.left_margin = Inches(0.69)
        ns.right_margin = Inches(0.69)
        sectPr = ns._sectPr

        # Ensure body section does NOT have different_first_page (titlePg)
        existing_titlePg = sectPr.find(qn('w:titlePg'))
        if existing_titlePg is not None:
            sectPr.remove(existing_titlePg)

        # Remove ALL header/footer references from body section
        # so it inherits section 0's default header/footer
        ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        for tag in (f'{{{ns_w}}}headerReference', f'{{{ns_w}}}footerReference'):
            for elem in list(sectPr.findall(tag)):
                sectPr.remove(elem)

        old = sectPr.find(qn('w:cols'))
        if old is not None:
            sectPr.remove(old)
        sectPr.append(parse_xml(f'<w:cols {nsdecls("w")} w:num="2" w:space="144"/>'))
        self._body_section = ns

        # Set section 0's DEFAULT header to the running title
        # (body section inherits this because it has no headerReference)
        header = self._title_section.header
        hp = header.paragraphs[0]
        hp.text = ""
        run = hp.add_run(running_title)
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        run.italic = True

        # Set section 0's DEFAULT footer to page number
        footer = self._title_section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.text = ""
        run = fp.add_run(f'VOLUME {volume}, 2025\t')
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        add_page_number(fp)

    def add_title(self, text):
        return self.doc.add_paragraph(text, style='PaperTitle')

    def add_author(self, text):
        return self.doc.add_paragraph(text, style='AU')

    def add_affiliation(self, text):
        return self.doc.add_paragraph(text, style='PI')

    def add_abstract(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run('Abstract\u2014')
        r.italic = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
        r2 = p.add_run(text)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(10)
        return p

    def add_keywords(self, text):
        return self.doc.add_paragraph(f'INDEX TERMS {text}.', style='IndexTerms')

    def add_h1(self, text):
        """Major heading: rendered in ALL CAPS."""
        return self.doc.add_paragraph(text.upper(), style='H1')

    def add_h2(self, text):
        """Subsection heading: rendered as-is (typically all caps or title case)."""
        return self.doc.add_paragraph(text, style='H2')

    def add_h3(self, text):
        return self.doc.add_paragraph(text, style='H3')

    def add_para(self, text):
        return self.doc.add_paragraph(text)

    def add_table_title(self, text):
        return self.doc.add_paragraph(text, style='TableTitle')

    def add_figure(self, path, caption, width=Inches(3.2)):
        if os.path.exists(path):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path, width=width)
        return self.doc.add_paragraph(caption, style='FigCaption')

    def add_ref(self, text):
        return self.doc.add_paragraph(text, style='RefStyle')

    def make_table(self, headers, rows):
        t = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            self._set_cell(t.rows[0].cells[i], h, bold=True)
        for r, row in enumerate(rows):
            for c, val in enumerate(row):
                self._set_cell(t.rows[r + 1].cells[c], str(val))
        return t

    def _set_cell(self, cell, text, bold=False, size=9):
        cell.text = str(text)
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(size)
            run.bold = bold
            run.font.name = "Times New Roman"

    def save(self, path: str):
        self.doc.save(path)
        return path
